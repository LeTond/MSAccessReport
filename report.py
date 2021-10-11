from docx import Document
import pandas as pd


class Report:
    def generate_week_document(self, dict_: dict, period: str) -> None:
        """
        Creates word file in witch it creates a table with information about category of studies,
        count of studies with contrast, without contrast and all studies for ever category and saving
        it in root (program directory) directory
        :param dict_: diction with information from access file
        :param period: date interval witch determined by user
        :return: None
        """
        document = Document()
        document.add_heading('Отчет Отделения МРТ', level=1)
        document.add_paragraph(f'за неделю {period}').bold = True

        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ' '
        hdr_cells[1].text = 'С контрастом'
        hdr_cells[2].text = 'Без контраста'
        hdr_cells[3].text = 'Всего'

        for x, y, z, h in self.create_week_table(dict_):
            row_cells = table.add_row().cells
            row_cells[0].text = str(x)
            row_cells[1].text = str(y)
            row_cells[2].text = str(z)
            row_cells[3].text = str(h)

        document.add_page_break()
        document.save(f'./Недельный отчёт за {period}.docx')

    @staticmethod
    def create_week_table(dict_: dict) -> tuple:
        """
        Создает таблицу для недельного отчета
        :param dict_: Это словарь, в котором посчитано количество исследований
        :return: таблицу
        """
        table_records = (
            ('Всего исследований', dict_['C контрастом']['Всего'], dict_['Без контраста']['Всего'],
             dict_['Всего']['Всего']),
            ('ВМП', dict_['C контрастом']['ВМП'], dict_['Без контраста']['ВМП'], dict_['Всего']['ВМП']),
            ('ОМС стац', dict_['C контрастом']['ОМС стац'], dict_['Без контраста']['ОМС стац'],
             dict_['Всего']['ОМС стац']),
            ('Платно', dict_['C контрастом']['Пл'], dict_['Без контраста']['Пл'], dict_['Всего']['Пл']),
            ('Сотрудники', dict_['C контрастом']['Сотр'], dict_['Без контраста']['Сотр'], dict_['Всего']['Сотр']),
            ('ДМС', dict_['C контрастом']['ДМС'], dict_['Без контраста']['ДМС'], dict_['Всего']['ДМС']),
            ('ОМС амб', dict_['C контрастом']['ОМС амб'], dict_['Без контраста']['ОМС амб'], dict_['Всего']['ОМС амб']),
            ('Клиническая апробация', dict_['C контрастом']['КА'], dict_['Без контраста']['КА'], dict_['Всего']['КА']),
            ('Грант', dict_['C контрастом']['Грант'], dict_['Без контраста']['Грант'], dict_['Всего']['Грант']),
            ('Наука', dict_['C контрастом']['Наука'], dict_['Без контраста']['Наука'], dict_['Всего']['Наука'])
        )
        return table_records

    def generate_month_document(self, dict_: dict, period: str) -> None:
        """
        Creates word file in witch it creates a table with information about category of studies,
        count of studies with contrast, without contrast and all studies for ever category and saving
        it in root (program directory) directory
        :param dict_: diction with information from access file
        :param period: date interval witch determined by user
        :return: None
        """
        document = Document()
        document.add_heading('Отчет Отделения МРТ', level=1)
        document.add_paragraph(f'за {period}').bold = True

        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ' '
        hdr_cells[1].text = 'Всего'
        hdr_cells[2].text = 'из них с внутривенным  контрастированием'
        hdr_cells[3].text = 'в подразделениях, оказывающих медицинскую помощь в амбулаторных условиях'
        hdr_cells[4].text = 'в условиях дневного стационара'

        for x, y, z, h, j in self.create_month_table(dict_):
            row_cells = table.add_row().cells
            row_cells[0].text = str(x)
            row_cells[1].text = str(y)
            row_cells[2].text = str(z)
            row_cells[3].text = str(h)
            row_cells[4].text = str(j)

        document.add_page_break()
        document.save(f'./Месячный отчёт за {period}.docx')

    @staticmethod
    def create_month_table(dict_: dict) -> tuple:
        """
        Создает таблицу для месячного отчета
        :param dict_: Это словарь, в котором посчитано количество исследований
        :return: таблицу
        """
        table_records = (
            ('Всего выполнено МРТ', dict_['Всего']['Всего МРТ'], dict_['C контрастом']['Всего МРТ'],
             dict_['В поликлинике']['Всего МРТ'], '-'),

            ('в том числе: сердечно-сосудистой системы', dict_['Всего']['ССС'], dict_['C контрастом']['ССС'],
             dict_['В поликлинике']['ССС'], '-'),

            ('легких и средостения', dict_['Всего']['Легкие и средостения'], dict_['C контрастом']['Легкие и средостения'],
             dict_['В поликлинике']['Легкие и средостения'], '-'),

            ('органов брюшной полости и забрюшинного пространства', dict_['Всего']['ОБП и ЗП'], dict_['C контрастом']['ОБП и ЗП'],
             dict_['В поликлинике']['ОБП и ЗП'], '-'),

            ('органов малого таза', dict_['Всего']['Малый таз'], dict_['C контрастом']['Малый таз'],
             dict_['В поликлинике']['Малый таз'], '-'),

            ('молочной железы', '-', '-', '-', '-'),

            ('головного мозга', dict_['Всего']['Голова'], dict_['C контрастом']['Голова'],
             dict_['В поликлинике']['Голова'], '-'),

            ('шейного отдела', dict_['Всего']['Шейный отдел'], dict_['C контрастом']['Шейный отдел'],
             dict_['В поликлинике']['Шейный отдел'], '-'),

            ('грудного отдела', dict_['Всего']['Грудной отдел'], dict_['C контрастом']['Грудной отдел'],
             dict_['В поликлинике']['Грудной отдел'], '-'),

            ('пояснично-крестцового отдела', dict_['Всего']['Пояснично-крестцовый'], dict_['C контрастом']['Пояснично-крестцовый'],
             dict_['В поликлинике']['Пояснично-крестцовый'], '-'),

            ('области “голова-шея”', '-', '-', '-', '-'),

            ('костей, суставов и мягких тканей', dict_['Всего']['Суставы'], dict_['C контрастом']['Суставы'],
             dict_['В поликлинике']['Суставы'], '-'),

            ('сосудов', dict_['Всего']['Сосуды'], dict_['C контрастом']['Сосуды'],
             dict_['В поликлинике']['Сосуды'], '-'),

            ('Плаценты', dict_['Всего']['Плацента'], dict_['C контрастом']['Плацента'],
             dict_['В поликлинике']['Плацента'], '-'),

            ('Плода', dict_['Всего']['Плода'], dict_['C контрастом']['Плода'],
             dict_['В поликлинике']['Плода'], '-'),

            ('прочих органов и систем (плод, плацента, ХСО, орбиты)', dict_['Всего']['Прочие'], dict_['C контрастом']['Прочие'],
             dict_['В поликлинике']['Прочие'], '-'),

            ('Интервенционные вмешательства под МРТ – контролем (из стр.1)', '-', '-', '-', '-'),
        )

        return table_records

    @staticmethod
    def create_quarter_table(dict_: dict) -> tuple:
        """
        Создает таблицу для квартального отчета
        :param dict_: Это словарь, в котором посчитано количество исследований
        :return: таблицу
        """
        table_records = (
            ('Всего', dict_['amb']['without contrast']['Всего'], dict_['amb']['with contrast']['Всего'],
             'Всего', dict_['inpatient']['without contrast']['Всего'], dict_['inpatient']['with contrast']['Всего']),

            ('ОМС', dict_['amb']['without contrast']['ОМС'], dict_['amb']['with contrast']['ОМС'],
             'ОМС', dict_['inpatient']['without contrast']['ОМС'], dict_['inpatient']['with contrast']['ОМС']),

            ('Платно', dict_['amb']['without contrast']['Платно'], dict_['amb']['with contrast']['Платно'],
             'Платно', dict_['inpatient']['without contrast']['Платно'],
             dict_['inpatient']['with contrast']['Платно']),

            ('ДМС', dict_['amb']['without contrast']['ДМС'], dict_['amb']['with contrast']['ДМС'],
             'ДМС', dict_['inpatient']['without contrast']['ДМС'], dict_['inpatient']['with contrast']['ДМС']),
        )

        return table_records

    def generate_quarter_document(self, dict_: dict, period: str) -> None:
        """
        Creates word file in witch it creates a table with information about category of studies,
        count of studies with contrast, without contrast and all studies for ever category and saving
        it in root (program directory) directory
        :param dict_: diction with information from access file
        :param period: date interval witch determined by user
        :return: None
        """
        document = Document()
        document.add_heading('Отчет Отделения МРТ', level=1)
        document.add_paragraph(f'за квартал {period}').bold = True

        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Амбулаторно'
        hdr_cells[1].text = 'Без контраста'
        hdr_cells[2].text = 'С контрастом'
        hdr_cells[3].text = 'Стационарно'
        hdr_cells[4].text = 'Без контраста'
        hdr_cells[5].text = 'С контрастом'

        for x, y, z, h, l, m in self.create_quarter_table(dict_):
            row_cells = table.add_row().cells
            row_cells[0].text = str(x)
            row_cells[1].text = str(y)
            row_cells[2].text = str(z)
            row_cells[3].text = str(h)
            row_cells[4].text = str(l)
            row_cells[5].text = str(m)

        document.add_page_break()
        document.save(f'./Квартальный отчёт за {period}.docx')

    @staticmethod
    def generate_year_document(dict_: dict, period: str) -> None:
        """
        Creates word file in witch it creates a table with information about category of studies,
        count of studies with contrast, without contrast and all studies for ever category and saving
        it in root (program directory) directory
        :param dict_: diction with information from access file
        :param period: date interval witch determined by user
        :return: None
        """
        df = pd.DataFrame(dict_)
        df.to_excel(f"./Годовой отчет за {period}.xlsx")
